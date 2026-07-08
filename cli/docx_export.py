"""
Standalone Microsoft Word (.docx) export (roadmap item 3, P1.2).

Deliberately independent of the Obsidian/Markdown-vault export
(cli/vault_export.py): a docx is a one-off artefact a user takes with them
(email it, attach it to a project report), not something meant to accumulate
in a synced vault, so this has no [export].vault_dir dependency and no
persistent output directory of its own -- the caller decides where the
rendered file ends up (see cli/web.py's POST /api/export/docx, which streams
it straight back as a download).

Mirrors vault_export.py's session-gating and content sections (same
EXPORTABLE_STATES, same summary/action-items/MoM sources) so the two exports
never silently diverge on "what counts as exportable" -- only the rendering
target differs.

Read-only with respect to data/todo.md and session state (same as
vault_export.py): this writes only to wherever the caller points it, never
into data/. Not token-gated for the same reason vault_export.py isn't --
it re-renders content a human already reviewed, it cannot touch todo.md.
Not registered as an MCP tool: this is a dashboard/CLI convenience export,
not something the LLM agent should be able to trigger arbitrarily -- the
agent-facing MCP surface stays at exactly 8 tools (mcp_server/server.py).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from cli.vault_export import EXPORTABLE_STATES, ExportError, _chained_session_ids, _session_date_and_slug
from mcp_server import state as state_mod
from mcp_server.schemas import validate_session_id
from mcp_server.todo import parse_todo


def export_session_docx(
    session_id: str,
    meetings_dir: Path | str,
    todo_path: Path | str,
    state_dir: Path | str,
    output_path: Path | str,
) -> Path:
    """Render one session into a .docx file at `output_path`. Returns
    `output_path` on success.

    Raises:
        ExportError: unknown session, or not yet reviewed-grade
        (PROPOSED/REVIEWED/APPLIED) -- same gate as vault_export.export_session.
    """
    validate_session_id(session_id)
    meetings_dir = Path(meetings_dir)
    output_path = Path(output_path)

    try:
        session = state_mod.load_session_state(state_dir, session_id)
    except FileNotFoundError as exc:
        raise ExportError(f"Unknown session '{session_id}'.") from exc
    if session.state not in EXPORTABLE_STATES:
        raise ExportError(
            f"Session '{session_id}' is {session.state.value}; only sessions that "
            "reached human review (PROPOSED/REVIEWED/APPLIED) are exportable."
        )

    date_str, slug = _session_date_and_slug(session_id)
    meeting_type = session.metadata.get("meeting_type", "general")

    summary_path = meetings_dir / f"{session_id}.summary.md"
    mom_path = meetings_dir / f"{session_id}.mom.md"

    items = [
        item for item in parse_todo(todo_path).items
        if item.session_id == session_id and item.status != "deleted"
    ]

    doc = Document()

    doc.add_heading(f"{slug} — {date_str}", level=0)

    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Date: {date_str}    Type: {meeting_type}    State: {session.state.value}").italic = True

    chained = _chained_session_ids(session_id, meetings_dir)
    if chained:
        doc.add_paragraph(f"Previous sessions: {', '.join(chained)}")

    if summary_path.exists():
        doc.add_heading("Summary", level=1)
        for line in summary_path.read_text(encoding="utf-8").strip().splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    if items:
        doc.add_heading("Action Items", level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        header_cells = table.rows[0].cells
        for cell, heading in zip(header_cells, ["Description", "Owner", "Owner Type", "Due", "Priority", "Status"]):
            cell.text = heading
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.description
            row_cells[1].text = item.owner or ""
            row_cells[2].text = item.owner_type or ""
            row_cells[3].text = item.due_date or ""
            row_cells[4].text = item.priority or ""
            row_cells[5].text = "done" if item.done else item.status

    if mom_path.exists():
        doc.add_heading("Minutes of Meeting", level=1)
        for line in mom_path.read_text(encoding="utf-8").strip().splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                doc.add_heading(stripped.lstrip("#").strip(), level=2)
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    footer = doc.add_paragraph()
    footer_run = footer.add_run("Generated by Meeting Agent — local & offline")
    footer_run.font.size = Pt(8)
    footer_run.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
