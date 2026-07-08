"""Tests for cli/docx_export.py (standalone Word export, roadmap item 3).

Mirrors tests/cli/test_vault_export.py's session-gating tests since both
exports share the same EXPORTABLE_STATES gate -- only the render target and
the destination-path semantics (no vault_dir requirement) differ here.
"""

from __future__ import annotations

import pytest
from docx import Document

from cli.docx_export import ExportError, export_session_docx
from mcp_server.state import State, create_session


def _seed(tmp_path, session_id: str, state: State, *, with_artifacts: bool = True):
    meetings = tmp_path / "meetings"
    meetings.mkdir(exist_ok=True)
    create_session(tmp_path / "state", session_id, tmp_path / ".lock", 1.0,
                   initial_state=state, meeting_type="is-call")
    if with_artifacts:
        (meetings / f"{session_id}.md").write_text("transcript text", encoding="utf-8")
        (meetings / f"{session_id}.summary.md").write_text("- decided X\n", encoding="utf-8")
        (meetings / f"{session_id}.mom.md").write_text("# MoM\ncontent line", encoding="utf-8")
    return meetings


@pytest.fixture()
def todo(tmp_path):
    todo_path = tmp_path / "todo.md"
    todo_path.write_text(
        '- [ ] Draft ADR <!-- meta: {"id": "a1", "session_id": "is-call-20260701-090000", '
        '"owner": "Naga", "owner_type": "self", "due_date": "2026-07-10", "priority": "HIGH"} -->\n'
        '- [x] Unrelated <!-- meta: {"id": "b2", "session_id": "other"} -->\n',
        encoding="utf-8",
    )
    return todo_path


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_export_session_docx_writes_title_summary_items_and_mom(tmp_path, todo):
    meetings = _seed(tmp_path, "is-call-20260701-090000", State.APPLIED)
    out_path = tmp_path / "export" / "is-call-20260701-090000.docx"

    result = export_session_docx(
        "is-call-20260701-090000", meetings, todo, tmp_path / "state", out_path
    )

    assert result == out_path
    assert out_path.exists()

    doc = Document(str(out_path))
    text = _all_text(doc)
    assert "is-call — 2026-07-01" in text
    assert "decided X" in text
    assert "Draft ADR" in text
    assert "Naga" in text
    assert "self" in text  # owner_type surfaced in the action-items table
    assert "MoM" in text
    assert "content line" in text
    assert "Unrelated" not in text  # other sessions' items excluded


def test_export_session_docx_refuses_unreviewed_session(tmp_path, todo):
    meetings = _seed(tmp_path, "raw-20260701-090000", State.TRANSCRIBED)
    out_path = tmp_path / "export" / "raw.docx"
    with pytest.raises(ExportError, match="human review"):
        export_session_docx("raw-20260701-090000", meetings, todo, tmp_path / "state", out_path)
    assert not out_path.exists()


def test_export_session_docx_unknown_session_raises(tmp_path, todo):
    meetings = tmp_path / "meetings"
    meetings.mkdir(exist_ok=True)
    out_path = tmp_path / "export" / "unknown.docx"
    with pytest.raises(ExportError, match="Unknown session"):
        export_session_docx("does-not-exist", meetings, todo, tmp_path / "state", out_path)


def test_export_session_docx_creates_output_directory(tmp_path, todo):
    """The caller (POST /api/export/docx) points this at a fresh tmp path
    each time -- the parent directory must not need to pre-exist."""
    meetings = _seed(tmp_path, "is-call-20260701-090000", State.APPLIED)
    out_path = tmp_path / "brand-new-dir" / "nested" / "out.docx"

    export_session_docx("is-call-20260701-090000", meetings, todo, tmp_path / "state", out_path)

    assert out_path.exists()
